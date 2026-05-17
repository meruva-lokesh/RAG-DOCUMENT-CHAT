"""
Script to generate sample PDF and DOCX files for testing.
Run this once to create example documents.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
import os


def create_sample_docx():
    """Create a sample DOCX file."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Python Programming Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph(
        'Python is a high-level, interpreted programming language known for its simplicity '
        'and readability. Created by Guido van Rossum and first released in 1991, Python has '
        'become one of the most popular programming languages in the world.'
    )
    
    # Key Features
    doc.add_heading('Key Features', level=1)
    doc.add_paragraph('Easy to Learn', style='List Bullet')
    doc.add_paragraph('Versatile and Powerful', style='List Bullet')
    doc.add_paragraph('Large Standard Library', style='List Bullet')
    doc.add_paragraph('Active Community', style='List Bullet')
    doc.add_paragraph('Cross-Platform', style='List Bullet')
    
    # Basic Syntax
    doc.add_heading('Basic Syntax', level=1)
    doc.add_paragraph(
        'Python uses indentation to define code blocks, making it visually clean and easy to read. '
        'Unlike many other languages, Python does not use curly braces or semicolons to denote '
        'code structure.'
    )
    
    doc.add_paragraph(
        'Variables in Python are dynamically typed, meaning you don\'t need to declare their type. '
        'The interpreter automatically determines the type based on the assigned value.'
    )
    
    # Data Types
    doc.add_heading('Common Data Types', level=1)
    doc.add_paragraph('Integers (int): Whole numbers like 42, -17, 0')
    doc.add_paragraph('Floats (float): Decimal numbers like 3.14, -0.001, 2.0')
    doc.add_paragraph('Strings (str): Text data like "Hello", \'Python\', """Multi-line"""')
    doc.add_paragraph('Booleans (bool): True or False values')
    doc.add_paragraph('Lists (list): Ordered, mutable collections [1, 2, 3]')
    doc.add_paragraph('Tuples (tuple): Ordered, immutable collections (1, 2, 3)')
    doc.add_paragraph('Dictionaries (dict): Key-value pairs {"name": "John", "age": 30}')
    
    # Applications
    doc.add_heading('Applications', level=1)
    doc.add_paragraph(
        'Python is used in a wide variety of domains including web development (Django, Flask), '
        'data science and machine learning (NumPy, Pandas, TensorFlow), automation and scripting, '
        'game development, desktop applications, and scientific computing.'
    )
    
    # Conclusion
    doc.add_heading('Getting Started', level=1)
    doc.add_paragraph(
        'To start learning Python, download it from python.org, install it on your system, '
        'and begin with simple programs. Practice is key to mastering any programming language. '
        'The Python community offers extensive documentation, tutorials, and support to help '
        'beginners get started.'
    )
    
    # Save
    output_path = os.path.join('examples', 'sample.docx')
    doc.save(output_path)
    print(f"✅ Created {output_path}")


def create_sample_pdf():
    """Create a sample PDF file."""
    output_path = os.path.join('examples', 'sample.pdf')
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    
    # Container for content
    story = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#1f77b4',
        spaceAfter=30,
        alignment=TA_CENTER
    )
    heading_style = styles['Heading2']
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )
    
    # Title
    story.append(Paragraph("Artificial Intelligence Overview", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Introduction
    story.append(Paragraph("What is Artificial Intelligence?", heading_style))
    story.append(Paragraph(
        "Artificial Intelligence (AI) refers to the simulation of human intelligence in machines "
        "that are programmed to think and learn like humans. The field of AI research was founded "
        "on the claim that human intelligence can be so precisely described that a machine can be "
        "made to simulate it. This includes learning, reasoning, problem-solving, perception, and "
        "language understanding.",
        body_style
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # History
    story.append(Paragraph("Brief History", heading_style))
    story.append(Paragraph(
        "The term 'artificial intelligence' was coined in 1956 by John McCarthy at the Dartmouth "
        "Conference. Since then, AI has experienced several waves of optimism followed by "
        "disappointment and loss of funding, known as 'AI winters'. However, starting in the early "
        "2000s, advances in machine learning, increased computational power, and the availability "
        "of large datasets have led to significant breakthroughs in AI capabilities.",
        body_style
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # Types
    story.append(Paragraph("Types of AI", heading_style))
    story.append(Paragraph(
        "<b>Narrow AI (Weak AI):</b> AI systems that are designed to handle specific tasks, such as "
        "voice assistants, recommendation systems, or image recognition. These systems operate under "
        "a limited set of constraints and cannot generalize beyond their designated tasks.",
        body_style
    ))
    story.append(Paragraph(
        "<b>General AI (Strong AI):</b> AI systems with generalized human cognitive abilities. When "
        "presented with an unfamiliar task, a strong AI system can find a solution without human "
        "intervention. This level of AI does not yet exist.",
        body_style
    ))
    story.append(Paragraph(
        "<b>Super AI:</b> AI that surpasses human intelligence and ability. This is a hypothetical "
        "concept and subject of much debate in the AI research community.",
        body_style
    ))
    story.append(Spacer(1, 0.3*inch))
    
    # Page break
    story.append(PageBreak())
    
    # Applications
    story.append(Paragraph("Current Applications", heading_style))
    story.append(Paragraph(
        "AI is being applied across numerous industries and domains:",
        body_style
    ))
    story.append(Paragraph(
        "<b>Healthcare:</b> AI algorithms can analyze medical images, predict patient outcomes, "
        "assist in drug discovery, and personalize treatment plans. Machine learning models can "
        "detect patterns in patient data that humans might miss.",
        body_style
    ))
    story.append(Paragraph(
        "<b>Finance:</b> Banks and financial institutions use AI for fraud detection, algorithmic "
        "trading, credit scoring, and customer service chatbots. AI can process vast amounts of "
        "transaction data in real-time to identify suspicious activities.",
        body_style
    ))
    story.append(Paragraph(
        "<b>Transportation:</b> Self-driving cars use AI to perceive their environment, make "
        "decisions, and navigate safely. AI is also used in traffic management, route optimization, "
        "and predictive maintenance for vehicles.",
        body_style
    ))
    story.append(Paragraph(
        "<b>Education:</b> AI-powered tutoring systems can provide personalized learning experiences, "
        "adapt to individual student needs, and provide instant feedback on assignments.",
        body_style
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # Challenges
    story.append(Paragraph("Challenges and Ethical Considerations", heading_style))
    story.append(Paragraph(
        "As AI becomes more prevalent, several challenges and ethical considerations have emerged. "
        "Bias in AI systems can perpetuate or amplify existing societal biases if the training data "
        "is not carefully curated. Privacy concerns arise from the vast amounts of personal data "
        "required for AI training. The potential displacement of jobs due to automation is another "
        "significant concern. Additionally, questions about AI accountability, transparency, and "
        "the potential misuse of AI technologies require ongoing attention from researchers, "
        "policymakers, and society as a whole.",
        body_style
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # Future
    story.append(Paragraph("The Future of AI", heading_style))
    story.append(Paragraph(
        "The future of AI holds immense potential. Researchers are working on making AI systems "
        "more explainable, fair, and aligned with human values. Advances in areas like natural "
        "language processing, computer vision, and reinforcement learning continue to push the "
        "boundaries of what AI can achieve. As AI becomes more integrated into our daily lives, "
        "it will be crucial to ensure that these systems are developed and deployed responsibly, "
        "with consideration for their societal impact.",
        body_style
    ))
    
    # Build PDF
    doc.build(story)
    print(f"✅ Created {output_path}")


if __name__ == "__main__":
    # Ensure examples directory exists
    os.makedirs('examples', exist_ok=True)
    
    print("📄 Generating sample documents...")
    
    try:
        create_sample_pdf()
    except Exception as e:
        print(f"❌ Error creating PDF: {e}")
    
    try:
        create_sample_docx()
    except Exception as e:
        print(f"❌ Error creating DOCX: {e}")
    
    print("✅ Done!")
